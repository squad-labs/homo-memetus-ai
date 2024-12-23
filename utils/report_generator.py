import warnings
import os
from dotenv import load_dotenv
from typing import Annotated, TypedDict, List, Dict, Sequence
from langgraph.graph import StateGraph, START, END
from langgraph.graph.message import add_messages
from langchain_core.messages import BaseMessage
from pydantic import  Field, create_model
from openai import OpenAI
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import PromptTemplate
from docx import Document
from langchain_core.messages import HumanMessage
from langchain_community.tools.tavily_search import TavilySearchResults

load_dotenv()

warnings.filterwarnings("ignore") 
os.environ['OPENAI_API_KEY'] = os.getenv("OPENAI_API_KEY")
os.environ['TAVILY_API_KEY'] = os.getenv("TAVILY_API_KEY")
os.environ['REPORT_MODEL_ID'] = os.getenv("REPORT_MODEL_ID")

OPENAI_API_KEY = os.environ["OPENAI_API_KEY"]
TAVILY_API_KEY = os.environ["TAVILY_API_KEY"]
REPORT_MODEL_ID = os.environ["REPORT_MODEL_ID"]

class ReportGenerator(TypedDict):
  messages: Annotated[Sequence[BaseMessage], add_messages]
  outline: Dict[str, str]
  current_section: int
  section_content: str
  total_sections: int
  full_report: List[Dict[str, str]]

graph_builder = StateGraph(ReportGenerator)

print(graph_builder)

def create_outline_model(section_count: int):
  fields = {f"section{i}": (str, Field(description=f"Title for section {i}")) for i in range(1, section_count + 1)}
  return create_model("DynamicOutline", **fields)

fields = {f"section{i}": (str, Field(description=f"Title for section {i}")) for i in range(1, 3 + 1)}

DynamicOutline = create_model("DynamicOutline", **fields)

DynamicOutline = create_outline_model(3)

llm = ChatOpenAI(model=f"{REPORT_MODEL_ID}")

def outline_generator(state: ReportGenerator):
  DynamicOutline = create_outline_model(state["total_sections"])
  outline_parser = JsonOutputParser(pydantic_object=DynamicOutline)

  outline_prompt = PromptTemplate(
      template="""
      Create an outline for a detailed report with exactly {section_count} main sections.
      {format_instructions}
      The topic is: {topic}
      """,
      input_variables=["section_count", "topic"],
      partial_variables={"format_instructions": outline_parser.get_format_instructions()},
  )
  
  chain = outline_prompt | llm | outline_parser
  
  outline = chain.invoke({
    "section_count": state["total_sections"], 
    "topic": state["messages"][-1].content
  })
  return {"outline": outline}


search = TavilySearchResults(max_results=3)

client = OpenAI()

def contents_writer(state: ReportGenerator):
  if "error" in state:
    return {"messages": [AIMessage(content=f"An error occurred: {state['error']}")]}
  
  if state["current_section"] > state["total_sections"]:
    return {"messages": [AIMessage(content="Report completed.")]}
      
  current_section_key = f"section{state['current_section']}"
  current_topic = state["outline"][current_section_key]
  search_results = search.invoke(current_topic)    
  
  previous_sections_content = []
  for i in range(1, state['current_section']):
    section_key = f"section{i}"
    if section_key in state["section_content"]:
      previous_sections_content.append(f"""
      Section {i}: 
      {state['outline'][section_key]}
      {state['section_content'][section_key]}
      """)
  
  previous_sections = "\n\n".join(previous_sections_content)
  
  section_prompt = PromptTemplate(
      template="""
      Write a detailed section for the topic: {topic}. 
      
      Use the following search results for information: {search_results}

      It must be a detailed section with statistics and information.
      Also, It always related with Crypto Currency and Blockchain.
      Especially, Meme coin of Solana ecosystem.
      
      Previous sections:
      {previous_sections}
      Write only the content for this section, 
      do not include any image prompts or suggestions.
      Detailed statistics or information is needed, 
      so you should include collected information from search result.""",
      input_variables=["topic", "search_results", "previous_sections"],
  )
  section_content = llm.invoke(section_prompt.format(
      topic=current_topic,
      search_results=search_results,
      previous_sections=previous_sections
  ))

  return {
      "section_content": section_content.content,
      "current_section": state["current_section"] + 1,
      "full_report": state["full_report"] + [{
          "title": current_topic,
          "content": section_content.content
      }]
  }

def report_generator(state: ReportGenerator):
  doc = Document()
  doc.add_heading(f"Report: {state['messages'][0].content}", 0)

  for section in state['full_report']:
      doc.add_heading(section['title'], level=1)
      doc.add_paragraph(section['content'])
      doc.add_page_break()

  filename = f"report_{state['messages'][0].content}.docx".replace(" ", "_")
  doc.save(filename)

  return {
    "messages": [AIMessage(content=f"Report finalized and saved as {filename}.")],
    "report_file": filename
  }

graph_builder.add_node('outline_generator', outline_generator)
graph_builder.add_node('contents_writer', contents_writer)
graph_builder.add_node('report_generator', report_generator)

graph_builder.add_edge(START, 'outline_generator')
graph_builder.add_edge('outline_generator', 'contents_writer')
graph_builder.add_edge('report_generator', END)

def should_continue_writing(state: ReportGenerator):
  if state["current_section"] <= state["total_sections"]:
      return "write_section"
  else:
      return "finalize_report"

graph_builder.add_conditional_edges(
    "contents_writer",
    should_continue_writing,
    {
      "write_section": "contents_writer",
      "finalize_report": "report_generator"
    }
)
graph = graph_builder.compile()

current_directory = os.getcwd()

print(f"Current working path: {current_directory}")

topic = input("Write a topic of generating report: ")

total_sections = int(input("Write section count: "))

initial_state: ReportGenerator = {
  "messages": [HumanMessage(content=topic)],
  "total_sections": total_sections,
  "current_section": 1,
  "full_report": [],
}

for chunk in graph.stream(initial_state,stream_mode="update"):
    print(chunk)

print("\n=== Your report successfully generated ===")